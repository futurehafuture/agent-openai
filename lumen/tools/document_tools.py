"""Document & PDF tools: read, inspect, extract, convert, and search.

Reads plain text, Markdown, PDF (via pypdf), and Word .docx (via python-docx).
Text extraction is the foundation — the agent itself does the summarising and
reasoning over the returned text.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from pypdf import PdfReader

from ..logging_setup import get_logger
from ..workspace import workspace
from ._format import human_bytes, md_table
from .registry import register_tool

logger = get_logger(__name__)

_TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".rst", ".log", ".csv"}


def _read_pdf(path: Path, pages: set[int] | None = None) -> str:
    reader = PdfReader(str(path))
    out = []
    for i, page in enumerate(reader.pages):
        if pages is None or i in pages:
            out.append(page.extract_text() or "")
    return "\n\n".join(out).strip()


def _read_docx(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs).strip()


def _read_any(path: Path, pages: set[int] | None = None) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _read_pdf(path, pages)
    if suffix == ".docx":
        return _read_docx(path)
    if suffix in _TEXT_SUFFIXES:
        return path.read_text(encoding="utf-8", errors="replace").strip()
    raise ValueError(f"Unsupported document type '{suffix}'. Supported: pdf, docx, txt, md.")


def _parse_pages(spec: str) -> set[int] | None:
    """Parse a 1-based page spec like '1-3,5' into 0-based indices, or None for all."""
    spec = spec.strip()
    if not spec:
        return None
    pages: set[int] = set()
    for part in spec.split(","):
        part = part.strip()
        if "-" in part:
            start, end = part.split("-", 1)
            pages.update(range(int(start) - 1, int(end)))
        elif part:
            pages.add(int(part) - 1)
    return pages


def read_document(path: str, max_chars: int = 6000) -> str:
    """Read a document's text content (PDF, Word .docx, or plain text/Markdown).

    Args:
        path: Path to the document.
        max_chars: Truncate the returned text to this many characters.
    """
    resolved = workspace.resolve_read(path)
    if not resolved.exists():
        raise FileNotFoundError(workspace.display(resolved))
    text = _read_any(resolved)
    if not text:
        return f"**{workspace.display(resolved)}** appears to contain no extractable text."
    truncated = text[:max_chars]
    note = "" if len(text) <= max_chars else f"\n\n_…truncated; {len(text):,} chars total._"
    return f"**{workspace.display(resolved)}**\n\n{truncated}{note}"


def document_info(path: str) -> str:
    """Report metadata about a document: type, size, pages or word count.

    Args:
        path: Path to the document.
    """
    resolved = workspace.resolve_read(path)
    if not resolved.exists():
        raise FileNotFoundError(workspace.display(resolved))
    suffix = resolved.suffix.lower()
    size = human_bytes(resolved.stat().st_size)
    rows = [["path", workspace.display(resolved)], ["type", suffix or "?"], ["size", size]]
    if suffix == ".pdf":
        reader = PdfReader(str(resolved))
        rows.append(["pages", str(len(reader.pages))])
        meta = reader.metadata or {}
        if meta.get("/Title"):
            rows.append(["title", str(meta.get("/Title"))])
        if meta.get("/Author"):
            rows.append(["author", str(meta.get("/Author"))])
    else:
        text = _read_any(resolved)
        rows.append(["words", f"{len(text.split()):,}"])
        rows.append(["characters", f"{len(text):,}"])
    return md_table(["field", "value"], rows)


def extract_pdf_pages(path: str, pages: str = "", max_chars: int = 8000) -> str:
    """Extract text from specific PDF pages.

    Args:
        path: Path to a PDF file.
        pages: 1-based page selection, e.g. ``"1-3,5"``. Empty means all pages.
        max_chars: Truncate the returned text to this many characters.
    """
    resolved = workspace.resolve_read(path)
    if resolved.suffix.lower() != ".pdf":
        raise ValueError("extract_pdf_pages only works on .pdf files.")
    if not resolved.exists():
        raise FileNotFoundError(workspace.display(resolved))
    text = _read_pdf(resolved, _parse_pages(pages))
    label = pages or "all"
    truncated = text[:max_chars]
    note = "" if len(text) <= max_chars else "\n\n_…truncated._"
    return f"**{workspace.display(resolved)}** (pages: {label})\n\n{truncated}{note}"


def convert_document(path: str, to: str, output_name: str | None = None) -> str:
    """Convert a document to plain text or Word .docx.

    Supported: pdf→txt, docx→txt, txt/md→docx. (Use for quick format changes.)

    Args:
        path: Path to the source document.
        to: Target format: 'txt' or 'docx'.
        output_name: Optional output filename.
    """
    fmt = to.lower().lstrip(".")
    if fmt not in {"txt", "docx"}:
        raise ValueError("`to` must be 'txt' or 'docx'.")
    resolved = workspace.resolve_read(path)
    if not resolved.exists():
        raise FileNotFoundError(workspace.display(resolved))

    if fmt == "txt":
        text = _read_any(resolved)
        target = workspace.unique_output(output_name or f"{resolved.stem}.txt")
        target.write_text(text, encoding="utf-8")
    else:  # docx
        if resolved.suffix.lower() not in _TEXT_SUFFIXES:
            raise ValueError("Converting to .docx is supported from text/markdown sources.")
        target = workspace.unique_output(output_name or f"{resolved.stem}.docx")
        doc = Document()
        for line in resolved.read_text(encoding="utf-8", errors="replace").splitlines():
            stripped = line.strip()
            if stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith(("- ", "* ")):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            else:
                doc.add_paragraph(line)
        doc.save(str(target))
    return f"🔄 Converted to {fmt} → **{workspace.display(target)}**."


def search_documents(path: str, query: str, max_results: int = 20) -> str:
    """Search a folder of documents for a phrase and report which files contain it.

    Args:
        path: A directory containing documents to search (recurses).
        query: The text to search for (case-insensitive).
        max_results: Maximum number of matching files to report.
    """
    directory = workspace.resolve_dir(path)
    needle = query.lower().strip()
    if not needle:
        raise ValueError("`query` must be non-empty.")
    hits: list[list[str]] = []
    for f in directory.rglob("*"):
        if not f.is_file() or f.name.startswith("."):
            continue
        if f.suffix.lower() not in (_TEXT_SUFFIXES | {".pdf", ".docx"}):
            continue
        try:
            text = _read_any(f)
        except Exception:  # noqa: BLE001 - skip unreadable files, keep searching
            continue
        count = text.lower().count(needle)
        if count:
            idx = text.lower().find(needle)
            snippet = text[max(0, idx - 40): idx + 60].replace("\n", " ").strip()
            hits.append([workspace.display(f), str(count), f"…{snippet}…"])
        if len(hits) >= max_results:
            break
    if not hits:
        return f"No documents under **{workspace.display(directory)}** contain {query!r}."
    return (
        f"**{len(hits)} file(s) contain {query!r}**\n\n"
        + md_table(["file", "hits", "context"], hits, max_width=60)
    )


for _fn in (
    read_document,
    document_info,
    extract_pdf_pages,
    convert_document,
    search_documents,
):
    register_tool(_fn, category="skills")
